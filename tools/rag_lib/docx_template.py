from __future__ import annotations

import copy
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Inches


EXPERIENCE_SLOTS = [
    {"header": 18, "role": 19, "bullets": [20, 21, 22, 23, 24]},
    {"header": 25, "role": 26, "bullets": [27, 28, 29, 30, 31]},
    {"header": 32, "role": 33, "bullets": [34, 35, 36, 37, 38]},
    {"header": 39, "role": 40, "bullets": [41, 42, 43, 44]},
    {"header": 45, "role": 46, "bullets": [47, 48, 49]},
]
LEADERSHIP_SLOT = {"header": 51, "role": 52, "bullets": [53]}
EDUCATION_LINES = [4, 5, 6, 7, 8]
SKILL_LINES = [11, 12, 13, 14, 15]
HEADER_LINES = {"name": 0, "contact": 1}


@dataclass
class GeneratedResumeData:
    name: str
    contact: str
    education_lines: List[str]
    skill_lines: List[str]
    experiences: List[Dict[str, List[str] | str]]
    leadership_header: str
    leadership_role: str
    leadership_bullets: List[str]


def load_template(template_path: Path) -> Document:
    return Document(str(template_path))


def extract_budgets(doc: Document) -> Dict[int, int]:
    budgets: Dict[int, int] = {}
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            budgets[idx] = max(24, len(text))
    return budgets


def _set_paragraph_text(doc: Document, index: int, text: str) -> None:
    p = doc.paragraphs[index]
    if not p.runs:
        p.add_run(text)
        return

    # Keep original run-level formatting by distributing new text across
    # existing runs in proportion to original run lengths.
    original_lengths = [len(r.text) for r in p.runs]
    total_orig = sum(original_lengths)

    if total_orig <= 0:
        for r in p.runs:
            r.text = ""
        p.runs[0].text = text
        return

    n = len(text)
    allocated = []
    used = 0
    for i, orig_len in enumerate(original_lengths):
        if i == len(original_lengths) - 1:
            allocated.append(max(0, n - used))
            break
        share = int(round((orig_len / total_orig) * n))
        # Avoid consuming the entire remainder too early.
        max_allowed = max(0, n - used - (len(original_lengths) - i - 1))
        share = min(max_allowed, max(0, share))
        allocated.append(share)
        used += share

    cursor = 0
    for run, chunk_len in zip(p.runs, allocated):
        run.text = text[cursor : cursor + chunk_len]
        cursor += chunk_len


def _first_run_style(paragraph) -> dict:
    style = {"bold": None, "italic": None, "underline": None, "font_name": None, "font_size": None}
    if paragraph.runs:
        r = paragraph.runs[0]
        style = {
            "bold": r.bold,
            "italic": r.italic,
            "underline": r.underline,
            "font_name": r.font.name,
            "font_size": r.font.size,
        }
    return style


def _clear_paragraph_content(paragraph) -> None:
    p_elm = paragraph._p
    for child in list(p_elm):
        if child.tag.endswith("}pPr"):
            continue
        p_elm.remove(child)


def _add_run_with_style(paragraph, text: str, base_style: dict, *, bold=None) -> None:
    run = paragraph.add_run(text)
    run.bold = base_style.get("bold") if bold is None else bold
    run.italic = base_style.get("italic")
    run.underline = base_style.get("underline")
    if base_style.get("font_name"):
        run.font.name = base_style["font_name"]
    if base_style.get("font_size"):
        run.font.size = base_style["font_size"]


def _split_two_col(text: str) -> Tuple[str, str]:
    if "\t" not in text:
        return text.strip(), ""
    left, right = text.split("\t", 1)
    return left.strip(), right.strip()


def _rewrite_plain_paragraph(doc: Document, index: int, text: str, *, bold=None) -> None:
    p = doc.paragraphs[index]
    style = _first_run_style(p)
    _clear_paragraph_content(p)
    _add_run_with_style(p, text, style, bold=bold)


def _rewrite_labeled_line(doc: Document, index: int, text: str) -> None:
    p = doc.paragraphs[index]
    style = _first_run_style(p)
    _clear_paragraph_content(p)

    if ":" in text:
        label, rest = text.split(":", 1)
        _add_run_with_style(p, label, style, bold=True)
        _add_run_with_style(p, f":{rest}", style, bold=False)
    else:
        _add_run_with_style(p, text, style, bold=None)


def _rewrite_two_col_paragraph(
    doc: Document,
    index: int,
    text: str,
    *,
    left_bold=None,
    right_bold=None,
    right_tab_inches: float = 7.40,
) -> None:
    p = doc.paragraphs[index]
    style = _first_run_style(p)
    _clear_paragraph_content(p)
    left, right = _split_two_col(text)
    _add_run_with_style(p, left, style, bold=left_bold)
    if right:
        _set_right_tab_stop(p, pos_inches=right_tab_inches)
        _add_run_with_style(p, "\t", style, bold=None)
        _add_run_with_style(p, right, style, bold=right_bold)


def _clone_paragraph_properties(doc: Document, src_index: int, dst_index: int) -> None:
    src = doc.paragraphs[src_index]
    dst = doc.paragraphs[dst_index]
    src_ppr = src._p.pPr
    dst_p = dst._p
    if dst_p.pPr is not None:
        dst_p.remove(dst_p.pPr)
    if src_ppr is not None:
        dst_p.insert(0, copy.deepcopy(src_ppr))


def _clone_run_style(doc: Document, src_index: int, dst_index: int) -> None:
    src = doc.paragraphs[src_index]
    dst = doc.paragraphs[dst_index]
    if not src.runs or not dst.runs:
        return
    template = src.runs[0]
    for run in dst.runs:
        run.bold = template.bold
        run.italic = template.italic
        run.underline = template.underline
        if template.font.name:
            run.font.name = template.font.name
        if template.font.size:
            run.font.size = template.font.size


def _fit(text: str, limit: int) -> str:
    # Keep tab separators for two-column lines.
    if "\t" in text:
        parts = text.split("\t", 1)
        left = _fit(parts[0], max(8, int(limit * 0.62)))
        right = _fit(parts[1], max(6, int(limit * 0.32)))
        return f"{left}\t{right}"

    cleaned = " ".join(text.split())
    if len(cleaned) <= limit:
        return cleaned
    # First pass: drop parenthetical info.
    import re

    shrunk = re.sub(r"\s*\([^)]*\)", "", cleaned)
    shrunk = " ".join(shrunk.split())
    if len(shrunk) <= limit:
        return shrunk
    # Hard trim.
    if limit <= 3:
        return shrunk[:limit]
    return shrunk[: limit - 3].rstrip() + "..."


def compose_two_col(left: str, right: str, width: int = 112) -> str:
    left_clean = " ".join((left or "").split())
    right_clean = " ".join((right or "").split())
    if not right_clean:
        return left_clean
    # Use tab separator so right column can align by right tab stop.
    return f"{left_clean}\t{right_clean}"


def _set_right_tab_stop(paragraph, pos_inches: float = 7.40) -> None:
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(pos_inches), WD_TAB_ALIGNMENT.RIGHT)


def _rewrite_single_run_paragraph(doc: Document, index: int, text: str) -> None:
    p = doc.paragraphs[index]
    first_style = _first_run_style(p)
    _clear_paragraph_content(p)
    _add_run_with_style(p, text, first_style, bold=None)


def apply_generated_data(doc: Document, data: GeneratedResumeData) -> None:
    budgets = extract_budgets(doc)

    _set_paragraph_text(doc, HEADER_LINES["name"], _fit(data.name, budgets.get(0, 40)))
    # Contact line may contain legacy hyperlink nodes in template; rewrite as plain single run.
    # Keep full contact line whenever possible; avoid email truncation.
    _rewrite_single_run_paragraph(doc, HEADER_LINES["contact"], _fit(data.contact, 220))

    for i, idx in enumerate(EDUCATION_LINES):
        text = data.education_lines[i] if i < len(data.education_lines) else ""
        if not text:
            _set_paragraph_text(doc, idx, "")
            continue
        fitted = _fit(text, budgets.get(idx, 120))
        if i in (0, 2):
            _rewrite_two_col_paragraph(doc, idx, fitted, left_bold=True, right_bold=False)
        elif i in (1, 3):
            _rewrite_plain_paragraph(doc, idx, fitted, bold=False)
        else:
            _rewrite_plain_paragraph(doc, idx, fitted, bold=None)

    for i, idx in enumerate(SKILL_LINES):
        text = data.skill_lines[i] if i < len(data.skill_lines) else ""
        if text:
            _rewrite_labeled_line(doc, idx, _fit(text, budgets.get(idx, 130)))
        else:
            _set_paragraph_text(doc, idx, "")

    for slot_idx, slot in enumerate(EXPERIENCE_SLOTS):
        if slot_idx < len(data.experiences):
            exp = data.experiences[slot_idx]
            _rewrite_two_col_paragraph(
                doc,
                slot["header"],
                _fit(str(exp["header"]), budgets.get(slot["header"], 130)),
                left_bold=True,
                right_bold=True,
            )
            _rewrite_two_col_paragraph(
                doc,
                slot["role"],
                _fit(str(exp["role"]), budgets.get(slot["role"], 130)),
                left_bold=False,
                right_bold=False,
            )
            bullets = list(exp["bullets"])
            bullet_style = doc.paragraphs[slot["bullets"][0]].style
            for i, para_idx in enumerate(slot["bullets"]):
                txt = bullets[i] if i < len(bullets) else ""
                doc.paragraphs[para_idx].style = bullet_style
                _clone_paragraph_properties(doc, slot["bullets"][0], para_idx)
                _set_paragraph_text(doc, para_idx, _fit(txt, budgets.get(para_idx, 130)) if txt else "")
                if txt:
                    _clone_run_style(doc, slot["bullets"][0], para_idx)
        else:
            _set_paragraph_text(doc, slot["header"], "")
            _set_paragraph_text(doc, slot["role"], "")
            for para_idx in slot["bullets"]:
                _set_paragraph_text(doc, para_idx, "")

    if data.leadership_header:
        _rewrite_two_col_paragraph(
            doc,
            LEADERSHIP_SLOT["header"],
            _fit(data.leadership_header, budgets.get(LEADERSHIP_SLOT["header"], 140)),
            left_bold=True,
            right_bold=True,
        )
    else:
        _set_paragraph_text(doc, LEADERSHIP_SLOT["header"], "")

    if data.leadership_role:
        _rewrite_two_col_paragraph(
            doc,
            LEADERSHIP_SLOT["role"],
            _fit(data.leadership_role, budgets.get(LEADERSHIP_SLOT["role"], 140)),
            left_bold=False,
            right_bold=False,
        )
    else:
        _set_paragraph_text(doc, LEADERSHIP_SLOT["role"], "")

    for i, para_idx in enumerate(LEADERSHIP_SLOT["bullets"]):
        txt = data.leadership_bullets[i] if i < len(data.leadership_bullets) else ""
        _set_paragraph_text(doc, para_idx, _fit(txt, budgets.get(para_idx, 150)) if txt else "")


def save_document(doc: Document, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
